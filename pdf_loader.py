import os
import tempfile
from pypdf import PdfReader
from docx import Document
from langchain.schema import Document as LangChainDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter

def load_pdfs(uploaded_files):
    documents = []
    for uploaded_file in uploaded_files:
        file_ext = os.path.splitext(uploaded_file.name)[1].lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
            temp_file.write(uploaded_file.getvalue())
            temp_path = temp_file.name

        try:
            if file_ext == ".pdf":
                reader = PdfReader(temp_path)
                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text() or ""
                    if text.strip():
                        documents.append(
                            LangChainDocument(
                                page_content=text,
                                metadata={"file_name": uploaded_file.name, "page": page_num + 1}
                            )
                        )
            elif file_ext in [".docx", ".doc"]:
                doc = Document(temp_path)
                full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                if full_text.strip():
                    documents.append(
                        LangChainDocument(
                            page_content=full_text,
                            metadata={"file_name": uploaded_file.name, "page": 1}
                        )
                    )
        except Exception as e:
            print(f"Error loading {uploaded_file.name}: {e}")
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)
    return documents

def split_documents(documents):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    return text_splitter.split_documents(documents)